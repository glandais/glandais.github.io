from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from scripts.resume_loader import get_filename_base

DIST_DIR = Path(__file__).resolve().parent.parent / "dist"

FONT_NAME = "Calibri"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_NAME = Pt(14)
FONT_SIZE_SECTION = Pt(12)
MARGIN = Cm(2)


def _set_font(run, size=FONT_SIZE_BODY, bold=False):
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold


def _add_heading_text(doc, text, size=FONT_SIZE_SECTION):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=size, bold=True)
    p.space_after = Pt(4)
    return p


def _add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run)
    p.space_after = Pt(2)
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    _set_font(run)
    p.space_after = Pt(1)
    return p


def _format_date(d):
    if d is None:
        return "Présent"
    return d


def generate_docx(data: dict) -> Path:
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
        section.different_first_page_header_footer = False

    basics = data["basics"]

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(basics["name"])
    _set_font(run, size=FONT_SIZE_NAME, bold=True)

    # Label
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(basics["label"])
    _set_font(run, size=FONT_SIZE_BODY)

    # Contact info — in body, NOT in header/footer
    contact_parts = []
    if basics.get("email"):
        contact_parts.append(basics["email"])
    if basics.get("phone"):
        contact_parts.append(basics["phone"])
    loc = basics.get("location", {})
    if loc.get("city"):
        city_str = loc["city"]
        if loc.get("postalCode"):
            city_str = f"{loc['postalCode']} {city_str}"
        contact_parts.append(city_str)
    if basics.get("linkedin"):
        contact_parts.append(basics["linkedin"])
    if basics.get("github"):
        contact_parts.append(basics["github"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(" | ".join(contact_parts))
    _set_font(run, size=Pt(10))
    p.space_after = Pt(8)

    # Summary
    if basics.get("summary"):
        _add_heading_text(doc, "Profil")
        _add_body(doc, basics["summary"].strip())

    # Skills
    if data.get("skills"):
        _add_heading_text(doc, "Compétences")
        for group in data["skills"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{group['category']} : ")
            _set_font(run, bold=True)
            run = p.add_run(", ".join(group["keywords"]))
            _set_font(run)
            p.space_after = Pt(2)

    # Experience
    if data.get("experience"):
        _add_heading_text(doc, "Expérience professionnelle")
        for exp in data["experience"]:
            start = _format_date(exp.get("startDate"))
            end = _format_date(exp.get("endDate"))
            title_line = exp["title"]
            if exp.get("freelance"):
                title_line += " (Freelance)"
            company = exp["company"]

            p = doc.add_paragraph()
            run = p.add_run(f"{title_line} — {company}")
            _set_font(run, bold=True)
            p.space_after = Pt(0)

            p = doc.add_paragraph()
            date_loc = f"{start} – {end}"
            if exp.get("location"):
                date_loc += f" | {exp['location']}"
            run = p.add_run(date_loc)
            _set_font(run, size=Pt(10))
            p.space_after = Pt(2)

            for h in exp.get("highlights", []):
                _add_bullet(doc, h)

            if exp.get("stack"):
                p = doc.add_paragraph()
                run = p.add_run("Stack : ")
                _set_font(run, bold=True, size=Pt(10))
                run = p.add_run(exp["stack"])
                _set_font(run, size=Pt(10))
                p.space_after = Pt(6)

    # Education
    if data.get("education"):
        _add_heading_text(doc, "Formation")
        for edu in data["education"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu['studyType']} {edu['area']} — {edu['institution']}")
            _set_font(run, bold=True)
            p.space_after = Pt(0)

            p = doc.add_paragraph()
            date_str = f"{edu.get('startDate', '')} – {edu.get('endDate', '')}"
            if edu.get("location"):
                date_str += f" | {edu['location']}"
            run = p.add_run(date_str)
            _set_font(run, size=Pt(10))
            p.space_after = Pt(4)

    # Projects
    if data.get("projects"):
        _add_heading_text(doc, "Projets personnels")
        for proj in data["projects"]:
            p = doc.add_paragraph()
            title = proj["name"]
            if proj.get("url"):
                title += f" ({proj['url']})"
            run = p.add_run(title)
            _set_font(run, bold=True)
            p.space_after = Pt(0)

            _add_body(doc, proj["description"])
            if proj.get("keywords"):
                p = doc.add_paragraph()
                run = p.add_run(", ".join(proj["keywords"]))
                _set_font(run, size=Pt(10))
                p.space_after = Pt(4)

    # Languages
    if data.get("languages"):
        _add_heading_text(doc, "Langues")
        for lang in data["languages"]:
            _add_body(doc, f"{lang['language']} — {lang['fluency']}")

    # Interests
    if data.get("interests"):
        _add_heading_text(doc, "Centres d'intérêt")
        _add_body(doc, ", ".join(data["interests"]))

    # Save
    out_dir = DIST_DIR / "docx"
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = get_filename_base(data) + ".docx"
    out_path = out_dir / filename
    doc.save(str(out_path))
    return out_path
