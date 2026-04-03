from pathlib import Path
from docx import Document
from scripts.resume_loader import load_resume
from scripts.generate_docx import generate_docx


def test_generate_docx_creates_file(tmp_path, monkeypatch):
    import scripts.generate_docx as mod
    monkeypatch.setattr(mod, "DIST_DIR", tmp_path)
    data = load_resume("fr")
    out = generate_docx(data)
    assert out.exists()
    assert out.suffix == ".docx"


def test_docx_contains_name(tmp_path, monkeypatch):
    import scripts.generate_docx as mod
    monkeypatch.setattr(mod, "DIST_DIR", tmp_path)
    data = load_resume("fr")
    out = generate_docx(data)
    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Gabriel Landais" in full_text


def test_docx_uses_calibri(tmp_path, monkeypatch):
    import scripts.generate_docx as mod
    monkeypatch.setattr(mod, "DIST_DIR", tmp_path)
    data = load_resume("fr")
    out = generate_docx(data)
    doc = Document(str(out))
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.name:
                assert run.font.name == "Calibri"


def test_docx_has_no_tables(tmp_path, monkeypatch):
    import scripts.generate_docx as mod
    monkeypatch.setattr(mod, "DIST_DIR", tmp_path)
    data = load_resume("fr")
    out = generate_docx(data)
    doc = Document(str(out))
    assert len(doc.tables) == 0


def test_docx_has_no_headers_footers(tmp_path, monkeypatch):
    import scripts.generate_docx as mod
    monkeypatch.setattr(mod, "DIST_DIR", tmp_path)
    data = load_resume("fr")
    out = generate_docx(data)
    doc = Document(str(out))
    for section in doc.sections:
        header_text = "".join(p.text for p in section.header.paragraphs)
        footer_text = "".join(p.text for p in section.footer.paragraphs)
        assert header_text.strip() == ""
        assert footer_text.strip() == ""
